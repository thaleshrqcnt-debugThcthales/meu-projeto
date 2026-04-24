"""
Exporta projeto cultural para DOCX usando python-docx.
"""
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.services.project_generator import SECTION_LABELS


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)


def _add_paragraph(doc: Document, text: str, bold: bool = False, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def _set_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def export_project_docx(project, sections: list, budgets: list, schedules: list, output_path: str) -> str:
    doc = Document()
    _set_margins(doc)

    # Estilo padrão
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # CAPA
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(project.public_title or project.internal_name or "Projeto Cultural")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = "Calibri"
    title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph()

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"Proponente: {project.proponent_name}\n"
        f"Módulo: {project.selected_module or 'não definido'}\n"
        f"Valor: R$ {project.intended_budget:,.2f}\n" if project.intended_budget else
        f"Proponente: {project.proponent_name}\n"
        f"Módulo: {project.selected_module or 'não definido'}\n"
    )
    meta_run.font.size = Pt(12)
    meta_run.font.name = "Calibri"

    city_state = f"{project.city or ''}/{project.state or ''}".strip("/")
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"{city_state} — {datetime.now().strftime('%B de %Y').capitalize()}")
    date_run.font.size = Pt(11)
    date_run.font.name = "Calibri"

    doc.add_page_break()

    # SEÇÕES DO PROJETO
    section_order = [
        "RESUMO", "APRESENTACAO", "JUSTIFICATIVA", "DIAGNOSTICO_TERRITORIAL",
        "OBJETIVO_GERAL", "OBJETIVOS_ESPECIFICOS", "PUBLICO_ALVO", "METAS",
        "METODOLOGIA", "ACOES_PREVISTAS", "CRONOGRAMA", "EQUIPE",
        "PLANO_ACESSIBILIDADE", "PLANO_COMUNICACAO", "CONTRAPARTIDA",
        "INDICADORES", "ORCAMENTO_NARRATIVO", "RISCOS_MITIGACAO",
        "RESULTADOS_ESPERADOS", "LEGADO", "CHECKLIST_ANEXOS", "ADEQUACAO_EDITAL",
    ]

    sections_dict = {s.section_name: s.content for s in sections if s.content}

    for sec_name in section_order:
        content = sections_dict.get(sec_name, "")
        label = SECTION_LABELS.get(sec_name, sec_name)
        _add_heading(doc, label, level=1)

        if content:
            for para in content.split("\n\n"):
                para = para.strip()
                if not para:
                    continue
                if para.startswith("| ") or para.startswith("|--"):
                    # Tabela Markdown simples
                    _add_markdown_table(doc, para)
                elif para.startswith("## "):
                    _add_heading(doc, para[3:], level=2)
                elif para.startswith("- ") or para.startswith("* "):
                    for item in para.split("\n"):
                        item = item.lstrip("- *").strip()
                        if item:
                            p = doc.add_paragraph(item, style="List Bullet")
                            p.runs[0].font.name = "Calibri"
                            p.runs[0].font.size = Pt(11)
                else:
                    _add_paragraph(doc, para)
        else:
            _add_paragraph(doc, "[Seção não preenchida]", bold=False)

        doc.add_paragraph()

    # ORÇAMENTO como tabela
    if budgets:
        doc.add_page_break()
        _add_heading(doc, "Orçamento Detalhado", level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        headers = ["Item", "Categoria", "Qtd", "Valor Unit.", "Total", "Justificativa"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        for b in budgets:
            row = table.add_row().cells
            row[0].text = b.item or ""
            row[1].text = b.category or ""
            row[2].text = str(b.quantity)
            row[3].text = f"R$ {b.unit_value:,.2f}"
            row[4].text = f"R$ {b.total_value:,.2f}"
            row[5].text = b.justification or ""

        total = sum(b.total_value for b in budgets)
        total_row = table.add_row().cells
        total_row[3].text = "TOTAL"
        total_row[3].paragraphs[0].runs[0].bold = True
        total_row[4].text = f"R$ {total:,.2f}"
        total_row[4].paragraphs[0].runs[0].bold = True

    # CRONOGRAMA como tabela
    if schedules:
        doc.add_page_break()
        _add_heading(doc, "Cronograma de Execução", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        headers = ["Fase", "Mês", "Atividade", "Responsável"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        for s in schedules:
            row = table.add_row().cells
            row[0].text = s.phase or ""
            row[1].text = s.month or ""
            row[2].text = s.activity or ""
            row[3].text = s.responsible or ""

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path)


def _add_markdown_table(doc: Document, md_table: str):
    """Converte tabela Markdown básica em tabela DOCX."""
    rows = [r for r in md_table.split("\n") if r.strip() and not r.strip().startswith("|--")]
    if not rows:
        return

    cols = [c.strip() for c in rows[0].strip("|").split("|")]
    table = doc.add_table(rows=len(rows), cols=len(cols))
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        cells = [c.strip() for c in row.strip("|").split("|")]
        for c_idx, cell_text in enumerate(cells):
            if c_idx < len(table.rows[r_idx].cells):
                table.rows[r_idx].cells[c_idx].text = cell_text
                if r_idx == 0:
                    try:
                        table.rows[r_idx].cells[c_idx].paragraphs[0].runs[0].bold = True
                    except IndexError:
                        pass


def export_audit_docx(audit_data: dict, project, output_path: str) -> str:
    doc = Document()
    _set_margins(doc)

    _add_heading(doc, "Relatório de Auditoria de Coerência", level=1)
    _add_paragraph(doc, f"Projeto: {project.public_title or project.internal_name}")
    _add_paragraph(doc, f"Proponente: {project.proponent_name}")
    _add_paragraph(doc, f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph()

    _add_heading(doc, "Status Geral", level=2)
    _add_paragraph(doc, audit_data.get("status", "N/A"), bold=True)
    _add_paragraph(doc, f"Score de coerência: {audit_data.get('score', 0)}/100")
    _add_paragraph(doc, audit_data.get("summary", ""))
    doc.add_paragraph()

    findings = audit_data.get("findings", [])
    for level_name, level_label in [("critico", "Erros Críticos"), ("medio", "Inconsistências Médias"), ("leve", "Melhorias Leves")]:
        level_findings = [f for f in findings if f.get("level") == level_name]
        if level_findings:
            _add_heading(doc, level_label, level=2)
            for f in level_findings:
                _add_paragraph(doc, f"[{f.get('category', '').upper()}] {f.get('location', '')}:", bold=True)
                _add_paragraph(doc, f"Problema: {f.get('problem', '')}")
                _add_paragraph(doc, f"Impacto: {f.get('impact', '')}")
                _add_paragraph(doc, f"Sugestão: {f.get('suggestion', '')}")
                doc.add_paragraph()

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path)


def export_score_docx(score_data: dict, project, output_path: str) -> str:
    doc = Document()
    _set_margins(doc)

    _add_heading(doc, "Parecer de Banca Simulada", level=1)
    _add_paragraph(doc, f"Projeto: {project.public_title or project.internal_name}")
    _add_paragraph(doc, f"Proponente: {project.proponent_name}")
    _add_paragraph(doc, f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph()

    _add_paragraph(
        doc,
        "AVISO: Esta é uma simulação técnica de preparação. "
        "A aprovação real depende da banca, concorrência e critérios definitivos do edital.",
        bold=True
    )
    doc.add_paragraph()

    _add_heading(doc, f"Nota Total: {score_data.get('total_score', 0)}/100", level=2)
    _add_paragraph(doc, f"Faixa competitiva: {score_data.get('competitive_range', 'N/A')}", bold=True)
    doc.add_paragraph()

    _add_heading(doc, "Notas por Critério", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Critério"
    hdr[1].text = "Nota"
    hdr[2].text = "Máximo"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    criterios = [
        ("C1 — Mérito cultural", "c1_score"),
        ("C2 — Adequação ao edital", "c2_score"),
        ("C3 — Viabilidade técnica", "c3_score"),
        ("C4 — Orçamento e cronograma", "c4_score"),
        ("C5 — Equipe, acessibilidade, impacto", "c5_score"),
    ]
    for label, key in criterios:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(score_data.get(key, 0))
        row[2].text = "20"

    doc.add_paragraph()

    _add_heading(doc, "Parecer", level=2)
    _add_paragraph(doc, score_data.get("opinion_text", ""))
    doc.add_paragraph()

    recs = score_data.get("recommendations", [])
    if recs:
        _add_heading(doc, "Recomendações", level=2)
        for rec in recs:
            p = doc.add_paragraph(str(rec), style="List Bullet")
            if p.runs:
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.size = Pt(11)

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path)
